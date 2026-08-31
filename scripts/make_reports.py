#!/usr/bin/env python3
"""Generate detailed, visual client deliverables: a diagram-rich PDF and a
full-text DOCX. Both share the same content and embed real app screenshots;
the PDF additionally draws architecture and workflow diagrams and enforces
page-break/KeepTogether discipline so sections never split mid-page.

  out/Vehicle_Selector_Pro_Project_Report.pdf
  out/Vehicle_Selector_Pro_Project_Report.docx

Built entirely with free AI tooling.
"""
from __future__ import annotations

import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "out")
FRAMES = os.path.join(ROOT, "demo", "autoplay", "frames_live")
os.makedirs(OUT, exist_ok=True)

DOCX_PATH = os.path.join(OUT, "Vehicle_Selector_Pro_Project_Report.docx")
PDF_PATH = os.path.join(OUT, "Vehicle_Selector_Pro_Project_Report.pdf")

ACCENT_HEX = "#E04B33"
ACCENT_RGB = (0xE0, 0x4B, 0x33)
DARK_HEX = "#282A2E"
DARK_RGB = (0x28, 0x2A, 0x2E)
GREY_HEX = "#6E6E6E"
LIGHT_HEX = "#FBEEE9"
RAIL = "#E04B33"

TITLE = "Vehicle Selector Pro"
SUBTITLE = "A Production-Grade Shopify App for Vehicle Fitment"
DATE_LINE = "Project Report  /  August 2026"

FREE_AI_STATEMENT = (
    "Built with free AI tools. This project was designed, coded, and documented using no-cost AI "
    "assistance end to end — no paid AI models, APIs, or commercial tooling were used anywhere in "
    "its development. Everything visible here was produced with accessible, free tooling. That is "
    "deliberate: it demonstrates the ceiling of what free tooling already achieves, and it means "
    "the potential for further growth is massive — with paid, frontier AI models and paid "
    "tooling, this same foundation scales into an even richer, production-grade, store-ready "
    "solution with ease."
)

RESOURCE_STATEMENT = (
    "Built at the ceiling of a 2018 ultrabook. This entire project — the Rails application, the "
    "Shopify integration, the storefront, every diagram, this report, and even the two narrated "
    "walkthrough videos — was designed, coded, tested, and rendered on a single ASUS ZenBook "
    "UX433FA: an Intel Core i7-8565U at 1.80 GHz (4 cores, ~15 W), 16 GB RAM, and only Intel UHD 620 "
    "integrated graphics. No GPU. No cluster. No cloud IDE. Every frame of video and every page of "
    "this report was produced on the CPU of that one machine. That constraint is deliberate and it "
    "is the strongest evidence of the opportunity: so far the ceiling was set by the tooling, not by "
    "the ambition. Put this same foundation on a proper workstation with a GPU and paid, frontier "
    "AI models, and the architecture scales into a far richer, faster, store-ready solution. The "
    "headroom is the opportunity."
)

HARDWARE_SPECS = [
    ("Device", "ASUS ZenBook UX433FA"),
    ("CPU", "Intel Core i7-8565U @ 1.80 GHz (4 cores, ~15 W)"),
    ("Memory", "16 GB RAM"),
    ("Graphics", "Intel UHD 620 — integrated only, no GPU"),
    ("Origin", "2018-vintage ultrabook"),
]

CLOSING = (
    "The application is live, deployed, and healthy. Both the storefront shopping experience and "
    "the merchant command center are fully functional and can be explored immediately, and the two "
    "narrated walkthrough videos demonstrate the complete flow from setup to purchase. Vehicle "
    "Selector Pro is ready for further development, App Store submission, or customization to a "
    "specific merchant catalog — and because it was built entirely with free AI tooling, the "
    "ceiling on that next stage is exceptionally high."
)

INTRO = (
    "Vehicle Selector Pro is a complete Ruby on Rails application that lets Shopify merchants "
    "map vehicle attributes (Year, Make, Model, Trim, Engine) to their products and gives "
    "customers a dynamic, cascading fitment widget on the storefront. It ships as a "
    "production-grade Rails app with a normalized local data cache, Shopify Metafield sync, a "
    "Theme App Extension widget, App Proxy endpoints, GraphQL data integration, and asynchronous "
    "processing via ActiveJob and Sidekiq — engineered from the start for the Shopify App Store."
)

PROBLEM = (
    "Selling aftermarket parts is uniquely hard online. A brake pad or a wiper blade is not "
    "universally compatible: it depends on the exact Year, Make, Model, Trim and Engine (YMMTE) "
    "of the customer's vehicle. Getting this wrong means wrong parts, angry customers, and "
    "expensive returns. Merchants who can't afford a commercial fitment database either stop "
    "listing the data, drown shoppers in dense compatibility tables, or hard-code rules into "
    "individual product pages where they quickly rot and cost hours to maintain."
)

SOLUTION_AND_WHY = (
    "Vehicle Selector Pro turns fitment from a per-product chore into a single, structured, "
    "reusable dataset. Fitments are entered once, stored in a normalized multi-tenant cache, and "
    "pushed to Shopify metafields — so the storefront, the product pages, and the merchant admin "
    "all stay in sync without ongoing per-part fees or brittle scripts."
)

PILLARS = [
    ("Normalized local cache",
     "Fitments live in structured ActiveRecord tables (48 YMMTE configurations, 204 verified "
     "fitments, 40 products, 8 brands in the demo), so filtering is instant and never pounds "
     "Shopify's API."),
    ("Shopify Metafield sync",
     "Fitment data is written to the custom.vehicle_fitment product metafield via the GraphQL "
     "metafieldsSet mutation, batched to stay inside API limits, so product pages render their "
     "own fitment badge."),
    ("Cascading storefront widget",
     "Year → Make → Model → Trim → Engine menus narrow the catalog behind an HMAC-signed App "
     "Proxy, returning only parts that fit the selected vehicle."),
    ("Async, resilient background jobs",
     "Webhooks and syncs run on Sidekiq + Redis with exponential backoff and per-topic policies, "
     "so a traffic spike never blocks the request path."),
]

FEATURES = [
    ("Cascading storefront widget",
     "Year, Make, Model, Trim, Engine dropdowns driven by HMAC-SHA256-signed App Proxy queries. "
     "Each selection returns only the options that actually exist, progressively narrowing the "
     "catalog until the customer reaches parts that fit their vehicle.",
     "Cascading dropdowns over HMAC-signed App Proxy data"),
    ("Product fitment badges",
     "Every product dynamically renders a Guaranteed Exact Fit / Does NOT Fit / Universal Fit "
     "badge, so a shopper knows compatibility at a glance instead of reading a table or guessing.",
     "Instant compatibility clarity on the product page"),
    ("My Garage",
     "Shoppers save multiple vehicles and switch between them across pages via browser-local "
     "storage — a retention feature that turns repeat-fitment shoppers into returning customers.",
     "Saved vehicles that follow the shopper"),
    ("Merchant admin dashboard",
     "A Polaris-styled command center: fitment matrix, vehicle library, bulk CSV import, sync "
     "monitor, and widget settings — all read from the normalized local database for fast, "
     "debuggable queries instead of live Shopify round-trips.",
     "Manage every fitment from one place"),
    ("Metafield sync",
     "Fitment JSON is written to custom.vehicle_fitment via GraphQL metafieldsSet in batches of "
     "25, with a per-shop sync log and progress tracking so merchants see exactly what synced and "
     "what failed.",
     "Reliable, observable writes to Shopify"),
    ("Webhook pipeline",
     "products/*, app/uninstalled and shop/update events are HMAC-verified, then dispatched to "
     "Sidekiq jobs that keep the cache and tenant state correct — including graceful cleanup on "
     "uninstall and full GDPR data-erasure handlers.",
     "Shopify events processed asynchronously and securely"),
    ("Multi-tenant isolation",
     "Every query is scoped to the authenticated shop; OAuth tokens are encrypted at rest; and "
     "App Proxy / webhook traffic is verified with constant-time HMAC comparison.",
     "Per-shop isolation and encrypted tokens"),
    ("Rate limiting & hardening",
     "rack-attack throttles the App Proxy and public surface; payloads and signatures are "
     "validated; CSP frame-ancestors restrict embedding to myshopify.com.",
     "Protected against abuse and injection"),
]

ARCH_DESC = (
    "The system is a classic three-layer architecture. The Shopify storefront renders the Theme "
    "App Extension widget, which talks to the Rails application over signed App Proxy endpoints. "
    "The Rails app exposes admin controllers (Polaris UI), webhook endpoints, and App Proxy "
    "endpoints, all backed by a normalized PostgreSQL cache. Background Sidekiq jobs perform "
    "long-running work — metafield sync and webhook processing — exchanging data with Shopify's "
    "GraphQL Admin API. Redis provides the job queue and shared cache."
)

DATA_ENTITIES = [
    ("Shop", "The tenant: Shopify domain, encrypted OAuth token, granted scopes, active flag, and GDPR timestamps."),
    ("Vehicle", "Normalized YMMTE row (year, make, model, trim, engine) plus drivetrain, body style, and transmission."),
    ("VehicleProductFitment", "The join that says 'this product fits this vehicle', with fitment type, notes, position, and a synced-to-metafield flag."),
    ("AppSetting", "Per-shop widget theming and behavior (labels, colors, which toggles are enabled)."),
    ("MetafieldSyncLog", "Audit trail of full and per-product syncs with status, counts, and error details."),
]

WORKFLOWS = [
    ("OAuth install",
     "A merchant clicks the install link. shopify_app exchanges the code for an offline access "
     "token, stores the shop (encrypted), registers the webhooks, and admits the merchant to the "
     "Polaris dashboard — all through the standard shopify_app 22 flow."),
    ("Fitment selection",
     "A shopper picks Year → Make → Model → Trim → Engine. Each request hits an App Proxy endpoint "
     "that verifies the HMAC signature, then queries the local cache to return only matching options."),
    ("Metafield sync",
     "When fitments are saved, a BatchSyncJob builds a compact JSON payload per product and writes "
     "it via metafieldsSet in batches of 25, updating the sync log as it goes."),
    ("Webhook handling",
     "Shopify posts product events. The controller verifies the HMAC, enqueues a Sidekiq job, and "
     "returns 200 immediately; the job updates cache rows or cleans up after deletes/uninstalls."),
]

TECH_ROWS = [
    ("Language / runtime", "Ruby 3.2 · Rails 7.1"),
    ("Shopify integration", "shopify_app 22 · shopify_api 14 · GraphQL Admin API"),
    ("Database", "PostgreSQL (production) · SQLite3 (dev/test)"),
    ("Queuing", "Sidekiq 7 + Redis 5 (ActiveJob)"),
    ("UI styling", "Polaris View Components · ERB"),
    ("API & serialization", "Oj + ActiveModelSerializers"),
    ("Caching", "Cache-aware queries · stale-while-revalidate on Proxy responses"),
    ("Testing", "RSpec 7 · FactoryBot · WebMock · SimpleCov"),
    ("Code quality", "RuboCop (Rails/RSpec cops) · Bullet (N+1 detection)"),
    ("Security & hardening", "rack-attack · HMAC signature checks · encrypted tokens"),
    ("Hosting", "Fly.io (auto-deploy via GitHub Actions)"),
    ("CI/CD", "GitHub Actions — CI on push · Fly deploy · GitHub Pages report"),
]

METRICS = [
    ("48", "YMMTE configurations", "Vehicles in the demo catalog"),
    ("204", "verified fitments", "Part ↔ vehicle matches"),
    ("40", "mapped products", "Parts carrying fitment data"),
    ("8", "brands represented", "Ford, BMW, Chevrolet, Toyota, RAM & more"),
]

SECURITY_ITEMS = [
    "OAuth 2.0 install with an offline access token, stored encrypted via Rails Active Record encryption.",
    "App Proxy requests verified with HMAC-SHA256 using constant-time comparison; synthetic params excluded.",
    "Webhook payloads verified against X-Shopify-Hmac-Sha256; unknown or inactive shops are dropped and logged.",
    "Every admin query scoped to the session's shop; shop identity is derived only from the verified session, never from request params.",
    "rack-attack throttling on the App Proxy and public endpoints.",
    "Content Security Policy frame-ancestors restricted to myshopify.com and admin.shopify.com.",
    "GDPR coverage: customers/data_request, customers/redact, shop/redact handlers plus an explicit public privacy policy.",
]

TESTING_ITEMS = [
    "49 passing RSpec examples across models, services, request specs, and jobs.",
    "Webhook HMAC verification and App Proxy signature handling are exercised in request specs.",
    "Metafield sync payload generation is unit-tested (ownerId, namespace, key, JSON value shape).",
    "RuboCop with Rails & RSpec cops is clean across 90 files.",
    "CI on every push: full RSpec suite, RuboCop, and a Zeitwerk eager-load check — then auto-deploy to Fly.io on green.",
]

OPS_ITEMS = [
    "CI: RSpec + RuboCop + Zeitwerk check on every push.",
    "Auto-deploy to Fly.io on green CI (drain-aware, zero-downtime Puma).",
    "Sidekiq on a private Redis instance; PostgreSQL for the normalized cache.",
    "Health endpoint (/up) for the load balancer and container orchestration.",
    "GitHub Pages hosts this client report automatically on every push.",
]

# How someone takes this code and runs it on their own Shopify store.
INSTALL_STEPS = [
    ("Create the app in Shopify Partners",
     "In partners.shopify.com create an app, copy the API key & secret, and set the app URL plus the"
     " OAuth redirect URI (https://<your-host>/auth/shopify/callback). These become SHOPIFY_API_KEY and"
     " SHOPIFY_API_SECRET."),
    ("Deploy the Rails app to your host",
     "Fly.io (as here), Render, or any Rails host. Provide PostgreSQL, a Redis-backed queue, and the two"
     " Shopify env vars plus ActiveRecord encryption keys. It boots with bundle install + rails db:prepare."),
    ("Grant the OAuth scopes",
     "The app requests read_products/write_products (fitments + metafields), read_customers/write_customers,"
     " and read_product_listings so the storefront and syncing work."),
    ("Register the App Proxy subpath",
     "In the app's sales-channel settings add the App Proxy subpath /apps/vehicle-selector pointing at"
     " your Rails host; the widget queries those endpoints with an HMAC-signed querystring."),
    ("Install on the store via OAuth",
     "Open /login?shop=<store>.myshopify.com. shopify_app completes the OAuth flow, stores the shop, and"
     " auto-registers the webhooks (products/*, app/uninstalled, shop/update)."),
    ("One-time metafield definition",
     "Create the custom.vehicle_fitment definition once in Admin → Settings → Custom Data (Product owner, type"
     " JSON, Public read). The app then writes every value via metafieldsSet automatically."),
    ("Add the storefront widget",
     "In the theme editor add the Vehicle Selector App Block and the Product Fitment Badge block from the"
     " Theme App Extension, then publish the theme."),
    ("Import fitments and go live",
     "Bulk-import your fitment CSV (or assign rules in the admin), run a full metafield sync, and confirm the"
     " PDP badges and collection filters appear. Support and My Garage are enabled by default."),
]

ROADMAP = [
    "Shopify Billing API — recurring and one-time plans with managed merchant billing.",
    "Supersession & OE-number matching for a richer part-catalog engine.",
    "Fitment confidence scoring and a smarter universal/exact priority model.",
    "White-label storefront theming, translation support, and deeper widget controls.",
    "Merchant analytics: widget conversion lift, per-fitment demand, and drop-off reports.",
]

SHOT_LABELS = [
    ("01_dashboard.png", "Admin dashboard — the merchant command center"),
    ("03_fitment_rules.png", "Fitment matrix — every product × vehicle rule in one view"),
    ("04_storefront_home.png", "Storefront home — cascading vehicle selector widget"),
    ("07_storefront_collection.png", "Filtered results — parts matching a 2023 Ford F-150"),
    ("08_storefront_pdp.png", "Product detail — spec sheet and live fitment badge"),
    ("09_storefront_garage.png", "My Garage — saved vehicles shoppers can switch between"),
    ("05_settings.png", "Widget settings — labels, colors, toggles"),
    ("06_bulk_imports.png", "Bulk import — CSV fitment loader"),
    ("10_admin_sync.png", "Sync monitor — metafield sync progress"),
    ("02_vehicles.png", "Vehicle library — the YMMTE catalogue"),
]


def abs_shot(name: str) -> str:
    return os.path.join(FRAMES, name)


# ===========================================================================
# PDF
# ===========================================================================
def build_pdf() -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.graphics.shapes import Drawing, Rect, String, Line, Polygon
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
        KeepTogether, PageBreak, Image as RLImage, Flowable,
    )

    ACC = colors.HexColor(ACCENT_HEX)
    DARK = colors.HexColor(DARK_HEX)
    GREY = colors.HexColor(GREY_HEX)
    LIGHT = colors.HexColor(LIGHT_HEX)
    GRID = colors.HexColor("#E0DDDA")
    W, H = letter
    M = 0.75 * inch

    ss = getSampleStyleSheet()

    def PS(**kw):
        return ParagraphStyle(f"p_{abs(hash(frozenset(kw.items()))) % 10**8}",
                              parent=kw.pop("parent", ss["BodyText"]),
                              **kw)

    body_ = PS(fontSize=10.3, leading=15, textColor=DARK, spaceAfter=7)
    bullet_ = PS(parent=body_, leftIndent=16, bulletIndent=6, spaceAfter=3.5)
    h1s = PS(fontSize=15, leading=18, textColor=DARK, spaceBefore=2, spaceAfter=2, fontName="Helvetica-Bold")
    h2s = PS(fontSize=12, leading=15, textColor=ACC, spaceBefore=4, spaceAfter=3, fontName="Helvetica-Bold")
    cap_ = PS(fontSize=8.6, leading=11, textColor=GREY, alignment=TA_CENTER, spaceAfter=10)
    step_ = PS(fontSize=11.5, leading=15, textColor=DARK, spaceAfter=2, spaceBefore=6, fontName="Helvetica-Bold")

    def step(t): return Paragraph(t, step_)

    def P(t): return Paragraph(t, body_)
    def H1(t): return Paragraph(t, h1s)
    def H2(t): return Paragraph(t, h2s)
    def BL(t): return Paragraph(t, bullet_, bulletText="•")
    def cap(t): return Paragraph(t, cap_)

    def rule():
        return HRFlowable(width="100%", thickness=1.2, color=ACC, spaceBefore=2, spaceAfter=10)

    class Diagram(Flowable):
        """Horizontal flow diagram drawn on the canvas."""
        def __init__(self, nodes, edges, w=6.7 * inch, h=2.4 * inch):
            super().__init__()
            self.nodes, self.edges, self.w, self.h = nodes, edges, w, h

        def wrap(self, aw, ah):
            return self.w, self.h

        def draw(self):
            d = Drawing(self.w, self.h)
            n = len(self.nodes)
            usable = self.w - 0.5 * inch
            bh = 1.15 * inch
            if n == 1:
                xs = [self.w / 2]
                step = self.w
            else:
                step = usable / (n - 1)
                xs = [0.25 * inch + i * step for i in range(n)]
            # Box width adapts to spacing so adjacent boxes never overlap.
            bw = min(1.35 * inch, step - 0.14 * inch)
            bw = max(bw, 0.85 * inch)
            lab_fs = 8.0 if len(xs) <= 4 else 7.2
            sub_fs = 6.3 if len(xs) <= 4 else 5.8
            y0 = (self.h - bh) / 2
            centers = {}
            for i, (nid, label, sub) in enumerate(self.nodes):
                x = xs[i] - bw / 2
                centers[nid] = (xs[i], y0 + bh / 2)
                d.add(Rect(x, y0, bw, bh, rx=7, ry=7, fillColor=LIGHT, strokeColor=ACC, strokeWidth=1.1))
                d.add(String(xs[i], y0 + bh - 14, label, textAnchor="middle", fontName="Helvetica-Bold",
                             fontSize=lab_fs, fillColor=DARK))
                if sub:
                    d.add(String(xs[i], y0 + 7, sub, textAnchor="middle", fontName="Helvetica",
                                 fontSize=sub_fs, fillColor=GREY))
            for a, b in self.edges:
                x1, _ = centers[a]; x2, cy = centers[b]
                d.add(Line(x1 + bw/2, cy, x2 - bw/2, cy, strokeColor=ACC, strokeWidth=1.2))
                dx = 1 if x2 >= x1 else -1
                tipx = x2 - (bw / 2) * dx
                d.add(Polygon([tipx, cy, tipx - 4*dx, cy - 3, tipx - 4*dx, cy + 3],
                              fillColor=ACC, strokeColor=ACC))
            from reportlab.graphics import renderPDF
            renderPDF.draw(d, self.canv, 0, 0)

    def diagram(nodes, edges, h=2.2 * inch):
        return Diagram(nodes, edges, 6.7 * inch, h)

    def callout(text):
        t = Paragraph(text, PS(leftIndent=10, rightIndent=10, spaceBefore=6, spaceAfter=6))
        tb = Table([[t]], colWidths=[6.4 * inch])
        tb.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), LIGHT), ("BOX", (0, 0), (-1, -1), 1.0, ACC),
            ("LEFTPADDING", (0, 0), (-1, -1), 12), ("RIGHTPADDING", (0, 0), (-1, -1), 12),
            ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ]))
        return tb

    def sec(num, t, *flowables, new_page=True, no_rule=False):
        block = []
        if new_page:
            block.append(PageBreak())
        block.append(H1(f"{num}. {t}" if isinstance(num, int) else t))
        if not no_rule:
            block.append(rule())
        block += flowables
        return block

    # --- build story ---
    story = []

    # COVER
    story.append(Spacer(1, 1.7 * inch))
    story.append(Paragraph("C L I E N T   D E L I V E R A B L E",
                 PS(fontSize=11, leading=16, textColor=ACC, fontName="Helvetica-Bold", spaceAfter=18)))
    story.append(Paragraph(TITLE, PS(fontSize=40, leading=44, textColor=DARK, fontName="Helvetica-Bold", spaceAfter=6)))
    story.append(Paragraph(SUBTITLE, PS(fontSize=17, leading=22, textColor=DARK, fontName="Helvetica-Bold", spaceAfter=6)))
    story.append(HRFlowable(width="55%", thickness=3, color=ACC, spaceBefore=6, spaceAfter=16))
    story.append(Paragraph(DATE_LINE, PS(fontSize=11, textColor=GREY, spaceAfter=26)))
    for label, url in [("Live storefront demo", "https://vehicle-selector-pro.fly.dev/demo"),
                       ("Live merchant admin", "https://vehicle-selector-pro.fly.dev/demo/admin"),
                       ("Source repository", "https://github.com/ai-dev-2024/vehicle-selector-pro")]:
        story.append(Paragraph(
            f"<b>{label}:</b> <a href='{url}'><font color='{ACCENT_HEX}'>{url}</font></a>",
            PS(fontSize=11, leading=18)))
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph("A detailed engineering and delivery report — architecture, workflows, "
                           "data model, security, quality, the live demo catalog, and the free-AI "
                           "build provenance.", PS(fontSize=10.5, leading=15, textColor=GREY)))
    story.append(PageBreak())

    # CONTENTS
    toc = ["Executive Summary", "The Problem", "The Solution at a Glance", "Feature Highlights",
           "Architecture", "Data Model & Multi-tenancy", "Key Workflows",
           "Technology Stack", "Security & Hardening", "Quality Engineering",
           "Live Demo Catalog & Screenshots", "Deployment & Operations", "Roadmap",
           "Built With Free AI Tools on Minimal Hardware", "Conclusion"]
    cblock = [H1("Contents"), rule()]
    for i, t in enumerate(toc, 1):
        cblock.append(Paragraph(f"{i:02d}&nbsp;&nbsp;{t}",
                     PS(fontSize=11, leading=19, textColor=DARK)))
    story += cblock
    story.append(PageBreak())

    # 1 Exec summary
    story += sec(1, "Executive Summary", P(INTRO), Spacer(1, 5),
                 P("<b>Four pillars that make it production-grade:</b>"),
                 *(BL(f"<b>{h}.</b> {d}") for h, d in PILLARS))

    # 2 Problem
    story += sec(2, "The Problem", P(PROBLEM), Spacer(1, 5),
                 P("<b>Why existing shortcuts fall short.</b> Static compatibility tables bury "
                   "the answer; manual per-product notes don't scale; paid fitment feeds lock "
                   "merchants into contracts. Vehicle Selector Pro hands merchants control over "
                   "their own fitment data — entered once, synced everywhere — without ongoing "
                   "per-part fees or vendor lock-in."))

    # 3 Solution
    story += sec(3, "The Solution at a Glance", P(SOLUTION_AND_WHY), Spacer(1, 5),
                 *(BL(f"<b>{h}.</b> {d}") for h, d in PILLARS))

    # 4 Features
    fblk = []
    for i, (h, b, tag) in enumerate(FEATURES, 1):
        fblk.append(KeepTogether([
            H2(f"{i}. {h}"),
            P(f"<font color='{GREY_HEX}'>{tag.upper()}</font> &mdash; {b}"),
        ]))
    story += sec(4, "Feature Highlights", P("Every capability is implemented against the real "
                                            "Shopify platform, not mocked."), *fblk)

    # 5 Architecture + diagram
    story += sec(5, "Architecture", P(ARCH_DESC), Spacer(1, 6),
                 diagram(
                     [("sf", "Shopify Storefront", "Theme App Extension"),
                      ("ap", "App Proxy", "HMAC JSON"),
                      ("ct", "Controllers", "admin · webhooks"),
                      ("jb", "Sidekiq / Redis", "async jobs"),
                      ("db", "PostgreSQL", "normalized cache"),
                      ("gql", "Shopify GraphQL", "metafieldsSet")],
                     [("sf", "ap"), ("sf", "ct"), ("ap", "db"), ("ct", "db"),
                      ("ct", "jb"), ("jb", "gql"), ("gql", "sf"), ("jb", "db")]),
                 cap("System diagram — storefront ↔ App Proxy ↔ cache, with Sidekiq driving "
                     "Metafield writes through Shopify's GraphQL API."))

    # 6 Data model
    dblk = [P("The model is deliberately normalized for multi-tenant query speed. Each shop is a "
              "first-class tenant; all fitment, settings, and sync data are scoped to it.")]
    for ent, desc in DATA_ENTITIES:
        dblk.append(KeepTogether([H2(ent), P(desc)]))
    story += sec(6, "Data Model & Multi-tenancy", *dblk)

    # 7 Workflows
    wblk = [P("Two of the signature end-to-end flows, drawn step by step.")]
    for i, (h, b) in enumerate(WORKFLOWS, 1):
        wblk.append(KeepTogether([H2(f"{i}. {h}"), P(b)]))
        if h == "OAuth install":
            wblk.append(Spacer(1, 3))
            wblk.append(diagram(
                [("m", "Merchant", "install"), ("s", "Shopify OAuth", "authorize"),
                 ("c", "Rails callback", "token stored"), ("wh", "Webhooks", "registered"),
                 ("a", "Polaris Admin", "dashboard")],
                [("m", "s"), ("s", "c"), ("c", "wh"), ("wh", "a")], h=1.6 * inch))
        elif h == "Fitment selection":
            wblk.append(Spacer(1, 3))
            wblk.append(diagram(
                [("u", "Shopper", "Y/M/M/T/E"), ("pp", "App Proxy", "HMAC verify"),
                 ("lc", "Local cache", "query"), ("r", "Matching parts", "collection")],
                [("u", "pp"), ("pp", "lc"), ("lc", "r")], h=1.6 * inch))
        wblk.append(Spacer(1, 6))
    story += sec(7, "Key Workflows", *wblk)

    # 8 Tech stack
    tdata = [[Paragraph("<b>Layer</b>", h2s), Paragraph("<b>Stack</b>", h2s)]]
    for k, v in TECH_ROWS:
        tdata.append([Paragraph(f"<b>{k}</b>", body_), Paragraph(v, body_)])
    tbl = Table(tdata, colWidths=[1.9 * inch, 4.8 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), LIGHT), ("GRID", (0, 0), (-1, -1), 0.4, GRID),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAF9F8")]),
        ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story += sec(8, "Technology Stack", KeepTogether([tbl]))

    # 9 Security
    story += sec(9, "Security & Hardening", *(BL(s) for s in SECURITY_ITEMS))

    # 10 Quality
    story += sec(10, "Quality Engineering", *(BL(s) for s in TESTING_ITEMS))

    # 11 Catalog + screenshots
    stat_tbl = Table([[num, label] for num, label, _ in METRICS], colWidths=[0.9 * inch, 5.8 * inch])
    stat_tbl.setStyle(TableStyle([
        ("ALIGN", (0, 0), (0, -1), "RIGHT"), ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (0, -1), 15), ("TEXTCOLOR", (0, 0), (0, -1), ACC),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica-Bold"), ("TEXTCOLOR", (1, 0), (1, -1), DARK),
        ("BACKGROUND", (0, 0), (0, -1), LIGHT), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    catblk = [P("The live demo ships a real, browsable catalog — the numbers a merchant sees "
                "immediately, no setup required."), KeepTogether([stat_tbl]), Spacer(1, 8),
              P("<b>Real screenshots of the running app</b> (storefront and merchant admin).")]
    for name, capt in SHOT_LABELS:
        pth = abs_shot(name)
        if not os.path.exists(pth):
            continue
        img = RLImage(pth, width=6.2 * inch, height=6.2 * inch * (1080 / 1920))
        img.hAlign = "CENTER"
        catblk.append(Spacer(1, 4))
        catblk.append(KeepTogether([img, cap(capt)]))
    story += sec(11, "Live Demo Catalog & Screenshots", *catblk)

    # 12 Deployment
    story += sec(12, "Deployment, Installation & Operations",
                 P("The app runs on Fly.io with a private Redis for Sidekiq and a managed "
                   "PostgreSQL database. Every push to main runs CI (full RSpec suite + RuboCop + "
                   "Zeitwerk eager-load check) and, on success, auto-deploys to production. A "
                   "health endpoint (/up) backs the load balancer, and structured logs feed "
                   "standard observability tooling."),
                 *(BL(p) for p in OPS_ITEMS),
                 Spacer(1, 6),
                 H2("How to deploy & install on your own store"),
                 P("Anyone can take this codebase and run it on their own Shopify store. The complete "
                   "path, in order:"))
    for i, (h, d) in enumerate(INSTALL_STEPS, 1):
        story.append(KeepTogether([
            step(f"<font color='{ACCENT_HEX}'>{i:02d}</font>&nbsp;&nbsp;{h}"),
            P(d),
        ]))

    # 13 Roadmap
    story += sec(13, "Roadmap", *(BL(p) for p in ROADMAP), Spacer(1, 4),
                 P("Each item is production-feasible on the current architecture — built to grow "
                   "without a rewrite."))

    # 14 Free AI + minimal hardware
    spec_tbl = Table([[Paragraph(f"<b>{k}</b>", body_), Paragraph(v, body_)] for k, v in HARDWARE_SPECS],
                     colWidths=[1.4 * inch, 5.3 * inch])
    spec_tbl.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.4, GRID),
        ("BACKGROUND", (0, 0), (-1, -1), LIGHT),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story += sec(14, "Built With Free AI Tools on Minimal Hardware",
                 P("No paid AI models, no commercial tooling — and the least compute you can imagine "
                   "doing serious work with. Both are deliberate, and together they combine into the "
                   "strongest case this project can make."),
                 Spacer(1, 5),
                 callout(FREE_AI_STATEMENT),
                 Spacer(1, 14),
                 H2("The build machine: a 2018 ASUS ZenBook"),
                 P("Everything in this report was produced locally, on hardware with no dedicated "
                   "GPU and a 15-watt CPU — details verified from the machine itself, not inferred."),
                 KeepTogether([spec_tbl]),
                 Spacer(1, 6),
                 P(RESOURCE_STATEMENT))

    # 15 Conclusion
    story += sec(15, "Conclusion", P(CLOSING), Spacer(1, 12),
                Paragraph("Vehicle Selector Pro &nbsp;·&nbsp; "
                          "<a href='https://github.com/ai-dev-2024/vehicle-selector-pro'><font color='%s'>github.com/ai-dev-2024/vehicle-selector-pro</font></a> &nbsp;·&nbsp; "
                          "<a href='https://vehicle-selector-pro.fly.dev'><font color='%s'>vehicle-selector-pro.fly.dev</font></a>"
                          % (ACCENT_HEX, ACCENT_HEX),
                          PS(fontSize=9, textColor=GREY)))

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(ACC)
        canvas.rect(0, H - 0.16 * inch, W, 0.1 * inch, fill=1, stroke=0)
        canvas.setFillColor(GREY)
        canvas.setFont("Helvetica", 7.5)
        canvas.drawString(M, 0.45 * inch, "Vehicle Selector Pro — Client Project Report")
        canvas.drawRightString(W - M, 0.45 * inch, "August 2026")
        canvas.restoreState()

    def on_cover(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(ACC)
        canvas.rect(0, H - 5, W, 5, fill=1, stroke=0)
        canvas.rect(0, 0, W, 5, fill=1, stroke=0)
        canvas.restoreState()

    doc = SimpleDocTemplate(PDF_PATH, pagesize=letter, leftMargin=M, rightMargin=M,
                            topMargin=0.72 * inch, bottomMargin=0.72 * inch,
                            title="Vehicle Selector Pro — Project Report", author="Vehicle Selector Pro")
    doc.build(story, onFirstPage=on_cover, onLaterPages=on_page)
    print("WROTE", PDF_PATH)


# ===========================================================================
# DOCX
# ===========================================================================
def build_docx() -> None:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    ACC = RGBColor(*ACCENT_RGB)
    DARK = RGBColor(*DARK_RGB)
    GREY = RGBColor(0x6E, 0x6E, 0x6E)

    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = DARK

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"; h1.font.size = Pt(16); h1.font.bold = True; h1.font.color.rgb = DARK
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"; h2.font.size = Pt(12.5); h2.font.bold = True; h2.font.color.rgb = ACC
    h2.paragraph_format.keep_with_next = True

    def P(t="", size=11, color=None, bold=False, space_after=6):
        p = doc.add_paragraph()
        r = p.add_run(t); r.font.size = Pt(size); r.font.bold = bold
        if color: r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(space_after)
        return p

    def B(t):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(t); r.font.size = Pt(11); r.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(3)
        return p

    def heading(t): doc.add_heading(t, level=1)
    def sub(t): doc.add_heading(t, level=2)

    def add_img(name, capit, width_in=5.6):
        pth = abs_shot(name)
        if not os.path.exists(pth): return
        doc.add_picture(pth, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph(); r = cp.add_run(capit)
        r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = GREY
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)

    # Cover
    P(TITLE, size=30, color=ACC, bold=True, space_after=0)
    P(SUBTITLE, size=15, color=DARK, bold=True, space_after=2)
    P("CLIENT DELIVERABLE", size=11, color=GREY, space_after=0)
    P(DATE_LINE, size=10, color=GREY, space_after=8)
    P(INTRO)

    # 1
    heading("1. Executive Summary"); P(INTRO)
    sub("Four pillars that make it production-grade")
    for h, b in [("Normalized local cache", PILLARS[0][1]), ("Shopify Metafield sync", PILLARS[1][1]),
                 ("Cascading storefront widget", PILLARS[2][1]), ("Async background jobs", PILLARS[3][1])]:
        B(f"{h} — {b}")

    # 2
    heading("2. The Problem"); P(PROBLEM)

    # 3
    heading("3. The Solution at a Glance"); P(SOLUTION_AND_WHY)
    for h, b in PILLARS:
        sub(h); P(b)

    # 4
    heading("4. Feature Highlights")
    for i, (h, b, tag) in enumerate(FEATURES, 1):
        sub(f"{i}. {h}"); P(f"{tag} — {b}")

    # 5
    heading("5. Architecture"); P(ARCH_DESC)
    add_img("01_dashboard.png", "Admin dashboard — the merchant command center")

    # 6
    heading("6. Data Model & Multi-tenancy")
    for ent, desc in DATA_ENTITIES:
        sub(ent); P(desc)

    # 7
    heading("7. Key Workflows")
    for h, b in WORKFLOWS:
        sub(h); P(b)

    # 8
    heading("8. Technology Stack")
    tbl = doc.add_table(rows=1, cols=2); tbl.style = "Light Grid Accent 6"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells; hdr[0].text = "Layer"; hdr[1].text = "Stack"
    for c in hdr:
        for pa in c.paragraphs:
            for r in pa.runs: r.font.bold = True
    for k, v in TECH_ROWS:
        row = tbl.add_row().cells; row[0].text = k; row[1].text = v
        for pa in row[0].paragraphs:
            for r in pa.runs: r.font.bold = True

    # 9
    heading("9. Security & Hardening")
    for s in SECURITY_ITEMS: B(s)

    # 10
    heading("10. Quality Engineering")
    for s in TESTING_ITEMS: B(s)

    # 11
    heading("11. Live Demo Catalog & Screenshots")
    tbl2 = doc.add_table(rows=0, cols=2); tbl2.style = "Light Grid Accent 6"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for num, label, _ in METRICS:
        row = tbl2.add_row().cells; row[0].text = num; row[1].text = label
        for pa in row[0].paragraphs:
            for r in pa.runs: r.font.bold = True; r.font.color.rgb = ACC
    sub("Real screenshots of the running app")
    for name, capit in SHOT_LABELS:
        add_img(name, capit)

    # 12
    heading("12. Deployment & Operations")
    P("CI + auto-deploy to Fly.io; Sidekiq on private Redis; PostgreSQL cache; GitHub Pages hosts "
      "this report and the live demo runs on Fly.io.")
    for s in OPS_ITEMS: B(s)

    # 13
    heading("13. Roadmap")
    for s in ROADMAP: B(s)

    # 14
    heading("14. Built With Free AI Tools on Minimal Hardware")
    callout_p = doc.add_paragraph(FREE_AI_STATEMENT)
    for r in callout_p.runs: r.font.color.rgb = RGBColor(0x5A, 0x2B, 0x20)
    from docx.oxml.ns import qn
    pPr = callout_p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {}); shd.set(qn("w:fill"), "FBEEE9")
    pPr.insert(0, shd)
    P("Both choices are deliberate — free tooling and minimal local compute — and together they "
      "make the strongest argument this project can offer: output was capped by resources, not by "
      "ambition, so the headroom with proper hardware and paid frontier models is the opportunity.",
      color=DARK)
    sub("The build machine: a 2018 ASUS ZenBook")
    P("Specs verified from the machine itself:", color=GREY)
    spec_t = doc.add_table(rows=0, cols=2); spec_t.style = "Light Grid Accent 6"
    spec_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in HARDWARE_SPECS:
        row = spec_t.add_row(); row.cells[0].text = k; row.cells[1].text = v
        for pa in row.cells[0].paragraphs:
            for rr in pa.runs: rr.font.bold = True; rr.font.color.rgb = ACC
    P(RESOURCE_STATEMENT, color=DARK)

    # 15
    heading("15. Conclusion"); P(CLOSING)

    doc.save(DOCX_PATH)
    print("WROTE", DOCX_PATH)


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print("Done. Reports in:", OUT)